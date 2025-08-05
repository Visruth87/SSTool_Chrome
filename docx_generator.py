"""
DOCX Generator
Creates Word documents with screenshots and URL information.
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from PIL import Image


class DOCXGenerator:
    def __init__(self):
        self.max_image_width = Inches(6)  # Maximum image width in document
        self.max_image_height = Inches(8)  # Maximum image height in document
        
    def create_document(self, screenshots_data, output_file):
        """
        Create a DOCX document with screenshots
        
        Args:
            screenshots_data (list): List of tuples (name, url, screenshot_path)
            output_file (str): Path to save the DOCX file
        """
        try:
            # Create new document
            doc = Document()
            
            # Add title
            self._add_title(doc)
            
            # Add summary
            self._add_summary(doc, screenshots_data)
            
            # Add page break
            doc.add_page_break()
            
            # Add screenshots
            for i, (name, url, screenshot_path) in enumerate(screenshots_data):
                self._add_screenshot_section(doc, name, url, screenshot_path, i + 1)
                
                # Add page break between screenshots (except for the last one)
                if i < len(screenshots_data) - 1:
                    doc.add_page_break()
            
            # Save document
            doc.save(output_file)
            
            return True
            
        except Exception as e:
            raise Exception(f"Failed to create DOCX document: {str(e)}")
            
    def _add_title(self, doc):
        """Add title section to document"""
        title = doc.add_heading('URL Screenshots Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add timestamp
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        date_para = doc.add_paragraph(f'Generated on: {timestamp}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add empty line
        doc.add_paragraph('')
        
    def _add_summary(self, doc, screenshots_data):
        """Add summary section to document"""
        doc.add_heading('Summary', level=1)
        
        summary_para = doc.add_paragraph(f'This report contains screenshots of {len(screenshots_data)} websites.')
        
        # Add table of contents
        doc.add_heading('Websites Included:', level=2)
        
        for i, (name, url, _) in enumerate(screenshots_data):
            doc.add_paragraph(f'{i + 1}. {name} - {url}', style='List Number')
            
    def _add_screenshot_section(self, doc, name, url, screenshot_path, index):
        """
        Add a screenshot section to the document
        
        Args:
            doc: Document object
            name (str): Name of the website
            url (str): URL of the website
            screenshot_path (str): Path to screenshot file
            index (int): Section index
        """
        try:
            # Add section heading
            heading = doc.add_heading(f'{index}. {name}', level=1)
            
            # Add URL information
            url_para = doc.add_paragraph()
            url_para.add_run('URL: ').bold = True
            url_para.add_run(url)
            
            # Add timestamp
            timestamp = self._get_image_timestamp(screenshot_path)
            if timestamp:
                time_para = doc.add_paragraph()
                time_para.add_run('Screenshot taken: ').bold = True
                time_para.add_run(timestamp)
            
            # Add some space
            doc.add_paragraph('')
            
            # Add screenshot if file exists
            if os.path.exists(screenshot_path):
                self._add_image_to_doc(doc, screenshot_path, name)
            else:
                doc.add_paragraph(f'Screenshot file not found: {screenshot_path}')
                
        except Exception as e:
            # Add error message to document
            error_para = doc.add_paragraph()
            error_para.add_run('Error: ').bold = True
            error_para.add_run(f'Failed to add screenshot for {name}: {str(e)}')
            
    def _add_image_to_doc(self, doc, image_path, name):
        """
        Add image to document with proper sizing
        
        Args:
            doc: Document object
            image_path (str): Path to image file
            name (str): Name for alt text
        """
        try:
            # Get image dimensions
            with Image.open(image_path) as img:
                original_width, original_height = img.size
                
            # Calculate scaling to fit within max dimensions
            width_scale = self.max_image_width.inches / (original_width / 96)  # Assuming 96 DPI
            height_scale = self.max_image_height.inches / (original_height / 96)
            scale = min(width_scale, height_scale, 1.0)  # Don't upscale
            
            # Calculate final dimensions
            final_width = Inches(original_width / 96 * scale)
            final_height = Inches(original_height / 96 * scale)
            
            # Add image to document
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.add_picture(image_path, width=final_width, height=final_height)
            
            # Add caption
            caption_para = doc.add_paragraph(f'Screenshot of {name}')
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Make caption italic and smaller
            for run in caption_para.runs:
                run.italic = True
                run.font.size = Pt(10)
                
        except Exception as e:
            # Add error message
            error_para = doc.add_paragraph(f'Error loading image {image_path}: {str(e)}')
            
    def _get_image_timestamp(self, image_path):
        """
        Get timestamp from image file
        
        Args:
            image_path (str): Path to image file
            
        Returns:
            str: Formatted timestamp or None
        """
        try:
            if os.path.exists(image_path):
                # Get file modification time
                mtime = os.path.getmtime(image_path)
                return datetime.fromtimestamp(mtime).strftime("%Y-%m-%d %H:%M:%S")
        except Exception:
            pass
        return None
        
    def create_document_with_multiple_images(self, screenshots_data, output_file):
        """
        Create document where each URL can have multiple screenshots
        
        Args:
            screenshots_data (list): List of tuples (name, url, [screenshot_paths])
            output_file (str): Path to save the DOCX file
        """
        try:
            doc = Document()
            
            # Add title
            self._add_title(doc)
            
            # Add summary
            total_screenshots = sum(len(paths) if isinstance(paths, list) else 1 
                                  for _, _, paths in screenshots_data)
            
            doc.add_heading('Summary', level=1)
            summary_para = doc.add_paragraph(
                f'This report contains {total_screenshots} screenshots from {len(screenshots_data)} websites.'
            )
            
            # Add page break
            doc.add_page_break()
            
            # Add screenshots
            for i, (name, url, screenshot_paths) in enumerate(screenshots_data):
                # Handle both single path and list of paths
                if isinstance(screenshot_paths, str):
                    screenshot_paths = [screenshot_paths]
                    
                doc.add_heading(f'{i + 1}. {name}', level=1)
                
                # Add URL information
                url_para = doc.add_paragraph()
                url_para.add_run('URL: ').bold = True
                url_para.add_run(url)
                
                doc.add_paragraph('')
                
                # Add all screenshots for this URL
                for j, screenshot_path in enumerate(screenshot_paths):
                    if len(screenshot_paths) > 1:
                        doc.add_heading(f'Screenshot {j + 1}', level=2)
                        
                    if os.path.exists(screenshot_path):
                        self._add_image_to_doc(doc, screenshot_path, f"{name} - {j + 1}")
                    else:
                        doc.add_paragraph(f'Screenshot file not found: {screenshot_path}')
                        
                    if j < len(screenshot_paths) - 1:
                        doc.add_paragraph('')  # Space between images
                
                # Add page break between URLs (except for the last one)
                if i < len(screenshots_data) - 1:
                    doc.add_page_break()
            
            # Save document
            doc.save(output_file)
            
            return True
            
        except Exception as e:
            raise Exception(f"Failed to create DOCX document: {str(e)}")
            
    def create_comparison_document(self, screenshots_data, output_file):
        """
        Create document with side-by-side comparisons if multiple screenshots per URL
        
        Args:
            screenshots_data (list): List of tuples (name, url, screenshot_paths)
            output_file (str): Path to save the DOCX file
        """
        try:
            doc = Document()
            
            # Add title
            self._add_title(doc)
            
            # Process each URL
            for i, (name, url, screenshot_paths) in enumerate(screenshots_data):
                if isinstance(screenshot_paths, str):
                    screenshot_paths = [screenshot_paths]
                    
                doc.add_heading(f'{i + 1}. {name}', level=1)
                
                # Add URL information
                url_para = doc.add_paragraph()
                url_para.add_run('URL: ').bold = True
                url_para.add_run(url)
                
                doc.add_paragraph('')
                
                # If multiple screenshots, create a table for comparison
                if len(screenshot_paths) > 1:
                    table = doc.add_table(rows=2, cols=len(screenshot_paths))
                    
                    # Add headers
                    for j, _ in enumerate(screenshot_paths):
                        cell = table.cell(0, j)
                        cell.text = f'Screenshot {j + 1}'
                        
                    # Add images
                    for j, screenshot_path in enumerate(screenshot_paths):
                        cell = table.cell(1, j)
                        if os.path.exists(screenshot_path):
                            # Add smaller image for table
                            paragraph = cell.paragraphs[0]
                            run = paragraph.add_run()
                            run.add_picture(screenshot_path, width=Inches(2.5))
                        else:
                            cell.text = 'Image not found'
                            
                else:
                    # Single screenshot
                    screenshot_path = screenshot_paths[0]
                    if os.path.exists(screenshot_path):
                        self._add_image_to_doc(doc, screenshot_path, name)
                    else:
                        doc.add_paragraph(f'Screenshot file not found: {screenshot_path}')
                
                # Add page break between URLs (except for the last one)
                if i < len(screenshots_data) - 1:
                    doc.add_page_break()
            
            # Save document
            doc.save(output_file)
            
            return True
            
        except Exception as e:
            raise Exception(f"Failed to create comparison document: {str(e)}") 